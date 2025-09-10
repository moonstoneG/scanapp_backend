# test_fill_doc.py
# pip install python-docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Optional
import os

TEMPLATE = "证据先行登记保存批准书.docx"
OUTPUT   = "generated_local.docx"

# ======== 数据模型（示例）========
class Item:
    def __init__(self, name: str, unit: Optional[str], qty: int):
        self.name = name
        self.unit = unit or ""
        self.qty  = int(qty or 0)

class Payload:
    def __init__(self, bureau: str, suspect: str, behavior: str, items: List[Item]):
        self.bureau   = bureau
        self.suspect  = suspect
        self.behavior = behavior
        self.items    = items

# ======== 文档工具 ========
def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def simple_run_replace(paragraph, placeholder: str, value: str, underline: bool = False):
    # 先尝试仅替换单个 run，最大程度保留样式
    for r in paragraph.runs:
        if placeholder in r.text:
            r.text = r.text.replace(placeholder, value)
            if underline:
                r.font.underline = True
            return True
    # 退化：整段替换（如占位符被拆到多个 run）
    if placeholder in paragraph.text:
        txt = paragraph.text
        paragraph.clear()
        before, after = txt.split(placeholder, 1)
        r1 = paragraph.add_run(before)
        r2 = paragraph.add_run(value)
        if underline:
            r2.font.underline = True
        paragraph.add_run(after)
        return True
    return False

def replace_core_placeholders(doc: Document, payload: Payload,
                              underline_bureau=False, underline_suspect=True, underline_behavior=True):
    for p in iter_all_paragraphs(doc):
        simple_run_replace(p, "{{BUREAU}}",   payload.bureau,   underline_bureau)
        simple_run_replace(p, "{{SUSPECT}}",  payload.suspect,  underline_suspect)
        simple_run_replace(p, "{{BEHAVIOR}}", payload.behavior, underline_behavior)

def replace_total_placeholders(doc: Document, kinds: int, total_qty: int) -> bool:
    """替换 {{TOTAL_KIND}} / {{TOTAL_QTY}}；返回是否至少替换了一个。"""
    replaced_any = False
    for p in iter_all_paragraphs(doc):
        if simple_run_replace(p, "{{TOTAL_KIND}}", str(kinds), underline=False):
            replaced_any = True
        if simple_run_replace(p, "{{TOTAL_QTY}}", str(total_qty), underline=False):
            replaced_any = True
    return replaced_any

def find_six_col_table_and_region(doc: Document):
    tbl = None
    for t in doc.tables:
        if len(t.columns) == 6:
            tbl = t
            break
    if tbl is None:
        raise RuntimeError("模板中未找到 6 列的明细表")

    start_row = 1
    end_row = len(tbl.rows)
    for i, row in enumerate(tbl.rows):
        line = " ".join(c.text for c in row.cells)
        if ("共计" in line) or ("总计" in line):
            end_row = i
            break
    if end_row <= start_row:
        end_row = len(tbl.rows)
    return tbl, start_row, end_row

def set_cell_center(cell):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def fill_items_left_right(doc: Document, items: List[Item]):
    """
    将 items 顺序写入 6 列表：左三列(名/单位/数量) -> 右三列(名/单位/数量) -> 下一行
    """
    tbl, start, end = find_six_col_table_and_region(doc)
    max_slots = (end - start) * 2
    items = items[:max_slots]

    row_i = start
    side = 0  # 0=左半区, 1=右半区
    for it in items:
        if row_i >= end:
            break
        row = tbl.rows[row_i]
        base = 0 if side == 0 else 3

        row.cells[base + 0].text = it.name
        row.cells[base + 1].text = it.unit
        row.cells[base + 2].text = str(it.qty)

        set_cell_center(row.cells[base + 0])
        set_cell_center(row.cells[base + 1])
        set_cell_center(row.cells[base + 2])

        if side == 1:
            row_i += 1
        side = 1 - side

def append_totals_numbers(doc: Document, kinds: int, total_qty: int):
    """
    兜底：在包含“总计（品种）”“总计（数量）”的段落/单元格 **后面** 追加数字。
    """
    def append_after_keyword_in_paragraph(p, keyword, number):
        if keyword not in p.text:
            return False
        last_run = p.runs[-1] if p.runs else None
        run = p.add_run(f" {number}")
        if last_run is not None:
            try:
                run.bold = last_run.bold
                run.italic = last_run.italic
                if last_run.font.size:
                    run.font.size = last_run.font.size
                if last_run.font.name:
                    run.font.name = last_run.font.name
            except Exception:
                pass
        return True

    wrote_kind = False
    wrote_qty  = False

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not wrote_kind and "总计（品种）" in p.text:
                        wrote_kind = append_after_keyword_in_paragraph(p, "总计（品种）", kinds)
                    if not wrote_qty and "总计（数量）" in p.text:
                        wrote_qty  = append_after_keyword_in_paragraph(p, "总计（数量）", total_qty)
                    if wrote_kind and wrote_qty:
                        break
            if wrote_kind and wrote_qty:
                break
        if wrote_kind and wrote_qty:
            break

    if not (wrote_kind and wrote_qty):
        for p in doc.paragraphs:
            if not wrote_kind and "总计（品种）" in p.text:
                wrote_kind = append_after_keyword_in_paragraph(p, "总计（品种）", kinds)
            if not wrote_qty and "总计（数量）" in p.text:
                wrote_qty  = append_after_keyword_in_paragraph(p, "总计（数量）", total_qty)
            if wrote_kind and wrote_qty:
                break

def generate_doc_local(payload: Payload,
                       underline_bureau=False, underline_suspect=True, underline_behavior=True,
                       template=TEMPLATE, output=OUTPUT):
    if not os.path.exists(template):
        raise FileNotFoundError(f"模板不存在：{template}")

    doc = Document(template)

    # 1) 核心占位符
    replace_core_placeholders(
        doc, payload,
        underline_bureau=underline_bureau,
        underline_suspect=underline_suspect,
        underline_behavior=underline_behavior
    )

    # 2) 表格填充（6 列）
    fill_items_left_right(doc, payload.items)

    # 3) 总计：先尝试占位符 {{TOTAL_KIND}} / {{TOTAL_QTY}}；若模板无这俩占位符则兜底追加
    kinds = len(payload.items)
    total_qty = sum(it.qty for it in payload.items)
    replaced = replace_total_placeholders(doc, kinds, total_qty)
    if not replaced:
        append_totals_numbers(doc, kinds, total_qty)

    doc.save(output)
    print(f"✅ 生成完成：{output}")

# ======== 本地示例运行 ========
if __name__ == "__main__":
    sample = Payload(
        bureau   = "桂平烟草局",
        suspect  = "张三",
        behavior = "无证经营烟草专卖品",
        items=[
            Item("中华", "条", 2),
            Item("玉溪", "盒", 5),
            Item("苏烟", "条", 1),
            Item("利群（软蓝）", "盒", 3),
            Item("黄鹤楼硬1916", "盒", 1),
        ],
    )
    # 若需要 {{BUREAU}} 也带下划线：underline_bureau=True
    generate_doc_local(sample, underline_bureau=False)