import cn2an
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

# ===== 数据模型 =====
class Item:
    def __init__(self, name: str, unit: str, qty: float, price: float):
        unit = (unit or "").strip()
        qty = float(qty)

        # 单位换算 → 条
        if unit in ["盒", "包"]:
            qty *= 0.1
        elif unit in ["箱", "件"]:
            qty *= 50

        self.name = name.strip()
        self.unit = "条"
        self.qty = round(qty, 1)
        self.price = float(price)

    @property
    def subtotal(self):
        return round(self.qty * self.price, 2)


class Payload:
    def __init__(self, bureau: str, items: list[Item]):
        self.bureau = bureau
        self.items = merge_items(items)


# ===== 工具函数 =====
def merge_items(items: list[Item]) -> list[Item]:
    """合并同名商品，数量累加，总价取加权平均（保留两位小数）"""
    merged = {}
    for it in items:
        if it.name not in merged:
            merged[it.name] = {"qty": 0.0, "subtotal": 0.0}
        merged[it.name]["qty"] += it.qty
        merged[it.name]["subtotal"] += it.subtotal

    result = []
    for name, data in merged.items():
        qty = round(data["qty"], 1)
        price = round(data["subtotal"] / qty, 2) if qty > 0 else 0.0
        # 这里不要再触发单位换算，直接 new 出来
        it = Item.__new__(Item)
        it.name, it.unit, it.qty, it.price = name, "条", qty, price
        result.append(it)
    return result


def _iter_blocks(doc):
    parent = doc
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _replace_in_blocks_after_table(doc: Document, tbl: Table, key: str, value: str):
    blocks = list(_iter_blocks(doc))
    try:
        i = blocks.index(tbl)
    except ValueError:
        return
    j = i + 1
    while j < len(blocks) and not isinstance(blocks[j], Table):
        blk = blocks[j]
        if isinstance(blk, Paragraph) and key in blk.text:
            for r in blk.runs:
                if key in r.text:
                    r.text = r.text.replace(key, value)
        j += 1


def _set_center(cell):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_page_placeholders_in_doc(doc: Document, total_pages: int):
    """
    替换正文里的 {{PAGE_INFO}} 占位符，按顺序填充页码
    """
    placeholders = []
    for p in doc.paragraphs:
        if "{{PAGE_INFO}}" in p.text:
            placeholders.append(p)

    for idx, p in enumerate(placeholders):
        cur_page = idx + 1
        text_val = f"第 {cur_page} 页 共 {total_pages} 页" if total_pages > 1 else ""
        # 清空 runs
        for r in p.runs:
            r.text = ""
        # 写新 run
        p.add_run(text_val)


# ===== 文档生成 =====
def generate_doc_pricing(payload, template="涉案物品核价表.docx", output="涉案物品核价表_生成.docx"):
    doc = Document(template)

    MAX_ROWS = 8
    tables = [t for t in doc.tables if len(t.columns) == 6]
    if not tables:
        raise RuntimeError("模板未找到 6 列表格")

    batches = [payload.items[i:i+MAX_ROWS] for i in range(0, len(payload.items), MAX_ROWS)]
    total_pages = len(batches)

    # 全局替换执法机关
    for p in doc.paragraphs:
        if "{{BUREAU}}" in p.text:
            for r in p.runs:
                if "{{BUREAU}}" in r.text:
                    r.text = r.text.replace("{{BUREAU}}", payload.bureau)

    used_tables = []
    for page_idx, (batch, tbl) in enumerate(zip(batches, tables)):
        used_tables.append(tbl)

        for row_idx in range(MAX_ROWS):
            row = tbl.rows[row_idx + 1]  # 第0行是表头
            if row_idx < len(batch):
                it = batch[row_idx]
                row.cells[0].text = str(page_idx * MAX_ROWS + row_idx + 1)
                row.cells[1].text = it.name
                row.cells[2].text = f"{it.qty:.1f}{it.unit}"
                row.cells[3].text = f"{it.price:.2f}"
                row.cells[4].text = f"{it.subtotal:.2f}"
                row.cells[5].text = ""
            else:
                for c in row.cells:
                    c.text = "/"

            for c in row.cells:
                _set_center(c)

    # === 合计只写在最后一张表的“合计”行 ===
    total_qty = sum(it.qty for it in payload.items)
    total_money = sum(it.subtotal for it in payload.items)

    if used_tables:
        last_tbl = used_tables[-1]
        for row in last_tbl.rows:
            first = row.cells[0].text.strip()
            if "合计" in first:
                row.cells[2].text = f"{total_qty:.1f} 条"
                row.cells[4].text = f"{total_money:.2f} 元"
                _set_center(row.cells[2])
                _set_center(row.cells[4])
                break

    # ✅ 页码替换
    replace_page_placeholders_in_doc(doc, total_pages)

    doc.save(output)
    print(f"✅ 已生成: {output}")
