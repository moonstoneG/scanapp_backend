# test_fill_doc.py
# pip install python-docx
import cn2an
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Optional
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATE = "证据先行登记保存批准书.docx"
OUTPUT   = "generated_local.docx"

# ======== 数据模型（示例）========
class Item:
    def __init__(self, name: str, unit: Optional[str], qty: float):
        self.name = name
        self.unit = "条"   # 统一为条
        qty = float(qty or 0.0)

        # 单位换算
        unit = (unit or "").strip()
        if unit == "盒":
            qty = qty * 0.1
        elif unit == "箱":
            qty = qty * 50
        # 其它情况默认当作「条」

        # ✅ 保留一位小数
        self.qty = round(qty, 1)

class Payload:
    def __init__(self, bureau: str, suspect: str, behavior: str, items: List[Item]):
        self.bureau   = bureau
        self.suspect  = suspect
        self.behavior = behavior
        self.items    = items

# ======== 文档工具 ========
def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def simple_run_replace(paragraph, placeholder: str, value: str, underline: bool = False):
    # 先尝试仅替换单个 run，最大程度保留样式
    for r in paragraph.runs:
        if placeholder in r.text:
            r.text = r.text.replace(placeholder, value)
            if underline:
                r.font.underline = True
            return True
    # 退化：整段替换（如占位符被拆到多个 run）
    if placeholder in paragraph.text:
        txt = paragraph.text
        paragraph.clear()
        before, after = txt.split(placeholder, 1)
        r1 = paragraph.add_run(before)
        r2 = paragraph.add_run(value)
        if underline:
            r2.font.underline = True
        paragraph.add_run(after)
        return True
    return False

def replace_core_placeholders(doc: Document, payload: Payload,
                              underline_bureau=False, underline_suspect=True, underline_behavior=True):
    for p in iter_all_paragraphs(doc):
        simple_run_replace(p, "{{BUREAU}}",   payload.bureau,   underline_bureau)
        simple_run_replace(p, "{{SUSPECT}}",  payload.suspect,  underline_suspect)
        simple_run_replace(p, "{{BEHAVIOR}}", payload.behavior, underline_behavior)

def replace_total_placeholders(doc: Document, kinds: int, total_qty: int) -> bool:
    """替换 {{TOTAL_KIND}} / {{TOTAL_QTY}}；返回是否至少替换了一个。"""
    replaced_any = False
    for p in iter_all_paragraphs(doc):
        if simple_run_replace(p, "{{TOTAL_KIND}}", str(kinds), underline=False):
            replaced_any = True
        if simple_run_replace(p, "{{TOTAL_QTY}}", str(total_qty), underline=False):
            replaced_any = True
    return replaced_any

def find_six_col_table_and_region(doc: Document):
    tbl = None
    for t in doc.tables:
        if len(t.columns) == 6:
            tbl = t
            break
    if tbl is None:
        raise RuntimeError("模板中未找到 6 列的明细表")

    start_row = 1  # 默认第 0 行是表头
    end_row = len(tbl.rows)

    for i, row in enumerate(tbl.rows):
        line = " ".join(c.text for c in row.cells)
        if ("共计" in line) or ("总计" in line):
            end_row = i   # 注意：总计行本身不算入可填充区
            break

    if end_row <= start_row:
        end_row = len(tbl.rows)

    return tbl, start_row, end_row

def set_cell_center(cell):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_diagonal(cell, direction="TL2BR"):
    """
    给单元格加斜线
    :param cell: cell 对象
    :param direction: "TL2BR" 左上到右下, "TR2BL" 右上到左下
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFFFF")
    tc_pr.append(shd)

    diag = OxmlElement("w:tcBorders")
    if direction == "TL2BR":
        br = OxmlElement("w:tl2br")
        br.set(qn("w:val"), "single")
        br.set(qn("w:sz"), "4")
        br.set(qn("w:space"), "0")
        br.set(qn("w:color"), "000000")
        diag.append(br)
    elif direction == "TR2BL":
        bl = OxmlElement("w:tr2bl")
        bl.set(qn("w:val"), "single")
        bl.set(qn("w:sz"), "4")
        bl.set(qn("w:space"), "0")
        bl.set(qn("w:color"), "000000")
        diag.append(bl)

    tc_pr.append(diag)

def fill_items_col_by_col(doc: Document, items: List[Item]):
    """
    先把左三列(名称/单位/数量)按行自上而下填满，再填右三列(名称/单位/数量)按行自上而下。
    """
    tbl, start, end = find_six_col_table_and_region(doc)
    nrows = max(end - start, 0)
    if nrows == 0 or not items:
        return

    # 左半区可容纳 nrows 条
    left_count = min(len(items), nrows)
    right_count = min(max(len(items) - left_count, 0), nrows)

    # 先填左三列
    for i in range(left_count):
        row = tbl.rows[start + i]
        it = items[i]
        row.cells[0].text = it.name
        row.cells[1].text = it.unit
        num_str = str(it.qty).rstrip("0").rstrip(".");
        row.cells[2].text = cn2an.an2cn(num_str, "up");
        set_cell_center(row.cells[0]); set_cell_center(row.cells[1]); set_cell_center(row.cells[2])

    # 再填右三列（从同一行起，自上而下）
    for i in range(right_count):
        row = tbl.rows[start + i]
        it = items[left_count + i]
        row.cells[3].text = it.name
        row.cells[4].text = it.unit
        num_str = str(it.qty).rstrip("0").rstrip(".");
        row.cells[5].text = cn2an.an2cn(num_str, "up");
        set_cell_center(row.cells[3]); set_cell_center(row.cells[4]); set_cell_center(row.cells[5])

def fill_items_left_right(doc: Document, items: List[Item]):
    """
    将 items 顺序写入 6 列表：左三列(名/单位/数量) -> 右三列(名/单位/数量) -> 下一行
    """
    def format_qty(qty):
            f = float(qty)
            return 
        
    tbl, start, end = find_six_col_table_and_region(doc)
    max_slots = (end - start) * 2
    items = items[:max_slots]

    row_i = start
    side = 0  # 0=左半区, 1=右半区
    for it in items:
        if row_i >= end:
            break
        row = tbl.rows[row_i]
        base = 0 if side == 0 else 3

        row.cells[base + 0].text = it.name
        row.cells[base + 1].text = it.unit
        num_str = str(it.qty).rstrip("0").rstrip(".");
        row.cells[base + 2].text = cn2an.an2cn(num_str, "up");

        set_cell_center(row.cells[base + 0])
        set_cell_center(row.cells[base + 1])
        set_cell_center(row.cells[base + 2])

        if side == 1:
            row_i += 1
        side = 1 - side

def append_totals_numbers(doc: Document, kinds: int, total_qty: int):
    """
    兜底：在包含“总计（品种）”“总计（数量）”的段落/单元格 **后面** 追加数字。
    """
    def append_after_keyword_in_paragraph(p, keyword, number):
        if keyword not in p.text:
            return False
        last_run = p.runs[-1] if p.runs else None
        run = p.add_run(f" {number}")
        if last_run is not None:
            try:
                run.bold = last_run.bold
                run.italic = last_run.italic
                if last_run.font.size:
                    run.font.size = last_run.font.size
                if last_run.font.name:
                    run.font.name = last_run.font.name
            except Exception:
                pass
        return True

    wrote_kind = False
    wrote_qty  = False

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not wrote_kind and "总计（品种）" in p.text:
                        wrote_kind = append_after_keyword_in_paragraph(p, "总计（品种）", kinds)
                    if not wrote_qty and "总计（数量）" in p.text:
                        wrote_qty  = append_after_keyword_in_paragraph(p, "总计（数量）", int(total_qty))
                    if wrote_kind and wrote_qty:
                        break
            if wrote_kind and wrote_qty:
                break
        if wrote_kind and wrote_qty:
            break

    if not (wrote_kind and wrote_qty):
        for p in doc.paragraphs:
            if not wrote_kind and "总计（品种）" in p.text:
                wrote_kind = append_after_keyword_in_paragraph(p, "总计（品种）", kinds)
            if not wrote_qty and "总计（数量）" in p.text:
                wrote_qty  = append_after_keyword_in_paragraph(p, "总计（数量）", int(total_qty))
            if wrote_kind and wrote_qty:
                break

def get_fill_region_for_table(tbl):
    """
    返回当前表格的可填充范围 [start_row, end_row)
    - 默认第0行是表头，从第1行开始写
    - 如果遇到“共计/总计”行，则该行不算入填充区域
    """
    start_row = 1
    end_row = len(tbl.rows)
    for i, row in enumerate(tbl.rows):
        line = "".join(c.text or "" for c in row.cells)
        if "共计" in line or "总计" in line:
            end_row = i  # 保留该行为总计，不写入数据
            break
    if end_row <= start_row:
        end_row = len(tbl.rows)
    return start_row, end_row

def replace_page_placeholder_in_table(tbl, cur_page: int, total_pages: int):
    """
    只替换当前表格内的 {{PAGE_INFO}} 占位符
    """
    text_val = f"第 {cur_page} 页 共 {total_pages} 页" if total_pages > 1 else ""
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                simple_run_replace(p, "{{PAGE_INFO}}", text_val, underline=False)

def generate_doc_local(payload: Payload,
                       underline_bureau=False, underline_suspect=True, underline_behavior=True,
                       template=TEMPLATE, output=OUTPUT):
    if not os.path.exists(template):
        raise FileNotFoundError(f"模板不存在：{template}")

    doc = Document(template)

    # ===== 替换所有页的占位符 =====
    replace_core_placeholders(
        doc, payload,
        underline_bureau=underline_bureau,
        underline_suspect=underline_suspect,
        underline_behavior=underline_behavior
    )

    # ===== 找所有 6 列表格 =====
    six_col_tables = [t for t in doc.tables if len(t.columns) == 6]
    if not six_col_tables:
        raise RuntimeError("模板中未找到 6 列明细表")

    # ===== 按 20 条一批分配到各个表格 =====
    MAX_ITEMS = 20
    batches = [payload.items[i:i+MAX_ITEMS] for i in range(0, len(payload.items), MAX_ITEMS)]
    total_pages = len(batches)
    for page_idx, (batch, tbl) in enumerate(zip(batches, six_col_tables)):
        cur_page = page_idx + 1
        start_row, end_row = get_fill_region_for_table(tbl)
        nrows = max(end_row - start_row, 0)
        if nrows <= 0:
         continue

        replace_page_placeholder_in_table(tbl, cur_page, total_pages)

        left_count = min(len(batch), nrows)
        right_count = min(max(len(batch) - left_count, 0), nrows)

        for i in range(left_count):
            row = tbl.rows[start_row + i]
            it = batch[i]
            row.cells[0].text = it.name
            row.cells[1].text = it.unit
            num_str = str(it.qty).rstrip("0").rstrip(".")
            row.cells[2].text = cn2an.an2cn(num_str, "up")
            set_cell_center(row.cells[0]); set_cell_center(row.cells[1]); set_cell_center(row.cells[2])

        for i in range(right_count):
            row = tbl.rows[start_row + i]
            it = batch[left_count + i]
            row.cells[3].text = it.name
            row.cells[4].text = it.unit
            num_str = str(it.qty).rstrip("0").rstrip(".")
            row.cells[5].text = cn2an.an2cn(num_str, "up")
            set_cell_center(row.cells[3]); set_cell_center(row.cells[4]); set_cell_center(row.cells[5])
        

    # ===== 总计（仍然只在第一页替换一次） =====
    kinds = len(payload.items)
    total_qty = sum(it.qty for it in payload.items)
    replaced = replace_total_placeholders(doc, kinds, float(total_qty))
    if not replaced:
        append_totals_numbers(doc, kinds, total_qty)

    doc.save(output)
    print(f"✅ 生成完成：{output}")
